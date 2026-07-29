# -*- coding: utf-8 -*-
"""関連研究に「音でデータを送る通信方式との関係」の項を足す。

2026-07-29、栗原さんがスプールを回しながら笛を吹いていて「モデムやファクシミリの音に
似ている」と気づいた。調べると似ているのは印象だけではなく、笛の符号化は周波数偏移変調
そのものであり、基準笛はパイロット信号、隣接同音禁止は自己同期のための符号、誤り訂正と
並べ替えは通信の定番と、一対一に対応する。この位置づけを関連研究に加える。

docx には生成器に無い栗原さんの手編集があるので、全体再生成はせず該当箇所だけを直す。
参考文献は末尾に [19] から [23] を足すだけにして、既存の番号は動かさない。
"""
import copy
import os
import shutil

from docx import Document
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "cipherflute_wiss2026_v1.2.docx")
PREV = os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_prev_20260729.docx")

ANCHOR_HEAD = "秘密の保管と鍵管理の実務"        # この見出しの直前へ差し込む
LAST_REF_HEAD = "[18] Shamir"

HEADING = "音でデータを送る通信方式との関係"

BODY1 = (
    "本研究の符号化は，音でデータを送る通信方式の系譜に連なる．電話回線のモデム[19]は"
    "2つの音の高さに0と1を割り当てる周波数偏移変調であり，ファクシミリの手順[20]や"
    "押しボタン信号[21]も，決められた音の組み合わせに意味を与えている．CipherFluteは"
    "12個の高さを使うので，その多値版に相当する．基準笛との比で読む仕組みは，受信側が"
    "周波数のずれを補正するために置く基準信号と同じ役割を果たす．隣り合う笛が同じ音に"
    "ならないようにする制約は，8B/10B[22]のように必ず遷移を起こして受信側が区切りを"
    "自力で見つけられるようにする符号と狙いが同じである．誤り訂正を付けて記号を並べ替える"
    "処理も，通信で広く使われてきたものである．空気を伝わる音でデータを送る試みそのものにも"
    "長い歴史があり[23]，音響カプラはその代表である．使う音域も1.7 kHzから3.1 kHzで，"
    "電話網が通す0.3 kHzから3.4 kHzの中に収まっている．あちらは通信路の制約，こちらは"
    "笛の大きさという別々の理由でありながら，結果として同じ帯に落ち着いている．")

BODY2 = (
    "違いは送信の側にある．通信では送信機が任意の信号をその場で作れるのに対し，"
    "CipherFluteでは物体の形そのものが信号を固定して保持する．電源も電子回路も要らない"
    "代わりに，位相や振幅は使えず，吹く速さも一定にならない．運べる量は毎秒10ビットに"
    "満たず，モデムには遠く及ばない．しかし本研究が求めるのは通信路の速度ではなく，"
    "電源なしで長期にわたり同じ信号を出し続ける担体である．速度と引き換えに保存性と"
    "自作可能性を得た点に，通信方式との位置づけの違いがある．")

REFS = [
    "[19] ITU-T: Recommendation V.21, 300 bits per second duplex modem standardized "
    "for use in the general switched telephone network (1988).",
    "[20] ITU-T: Recommendation T.30, Procedures for document facsimile transmission "
    "in the general switched telephone network (2005).",
    "[21] ITU-T: Recommendation Q.23, Technical features of push-button telephone sets (1988).",
    "[22] Widmer, A. X. and Franaszek, P. A.: A DC-Balanced, Partitioned-Block, 8B/10B "
    "Transmission Code, IBM Journal of Research and Development, Vol. 27, No. 5, "
    "pp. 440-451 (1983).",
    "[23] Madhavapeddy, A., Scott, D. and Sharp, R.: Audio Networking: The Forgotten "
    "Wireless Technology, IEEE Pervasive Computing, Vol. 4, No. 3, pp. 55-60 (2005).",
]


def set_text(p, txt):
    for r in list(p.runs)[1:]:
        r._element.getparent().remove(r._element)
    if p.runs:
        p.runs[0].text = txt
    else:
        p.add_run(txt)


def insert_after(p, txt):
    """p と同じ書式の段落を直後に作り、txt を入れて返す。"""
    el = copy.deepcopy(p._p)
    p._p.addnext(el)
    q = Paragraph(el, p._parent)
    set_text(q, txt)
    return q


def main():
    os.makedirs(os.path.join(HERE, "prev"), exist_ok=True)
    shutil.copy(SRC, PREV)
    d = Document(SRC)
    ps = d.paragraphs
    n_before = len(ps)

    anchor = next(p for p in ps if p.text.strip() == ANCHOR_HEAD)
    body_sample = next(p for p in ps if p.text.startswith("金属板に秘密を刻印する製品群"))

    # 見出しは anchor と同じ書式、本文は既存の本文段落と同じ書式で作る
    head_el = copy.deepcopy(anchor._p)
    anchor._p.addprevious(head_el)
    head = Paragraph(head_el, anchor._parent)
    set_text(head, HEADING)

    b1_el = copy.deepcopy(body_sample._p)
    head._p.addnext(b1_el)
    b1 = Paragraph(b1_el, head._parent)
    set_text(b1, BODY1)
    insert_after(b1, BODY2)

    # 参考文献は末尾に足す（既存の番号は動かさない）
    last = next(p for p in d.paragraphs if p.text.startswith(LAST_REF_HEAD))
    for ref in REFS:
        last = insert_after(last, ref)

    d.save(SRC)

    d2 = Document(SRC)
    t = [q.text for q in d2.paragraphs]
    print("段落数 %d → %d（+%d）" % (n_before, len(t), len(t) - n_before))
    print("新しい見出し:", HEADING in [x.strip() for x in t])
    print("参考文献[23]まで:", any(x.startswith("[23]") for x in t))
    print("概要は無改変:", t[4].startswith("概要．"))
    order = [i for i, x in enumerate(t) if x.strip() in (HEADING, ANCHOR_HEAD)]
    print("差し込み位置（新項→秘密の保管の順であるか）:", order, order == sorted(order))


if __name__ == "__main__":
    main()
