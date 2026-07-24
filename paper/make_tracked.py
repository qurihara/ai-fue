# -*- coding: utf-8 -*-
"""
「変更前 -> 変更後」の tracked changes（編集履歴）付き docx を作る補助スクリプト。

本論文の生成器（make_paper_wiss_v12.py）は毎回 docx を作り直す方式なので，素の
編集履歴は残らない。そこで生成器が prev/ に退避した「変更前」docx と，最新の
「変更後」docx を段落単位で突き合わせ，Word の編集履歴（w:ins / w:del）として
差分を書き込んだ docx を別ファイルとして出力する。

方針:
  - 本文の各段落を .text で並べ，difflib で「変更前」と「変更後」を対応づける。
  - 追加された段落は挿入（w:ins），削除された段落は削除（w:del）として記す。
  - 1対1で書き換えられた段落は，文字単位の差分を取り，変わった部分だけを
    インラインの挿入・削除として記す（日本語は分かち書きしないため文字単位）。
  - 図・表・空段落など本文テキストでない要素はそのまま残す。

LibreOffice の UNO/マクロによる比較は本環境では import uno が停止して不安定だった
ため，確実に動く python-docx による段落レベルの履歴付与を採用した。intra-段落は
1対1のときだけ文字単位に細分し，そうでない差分は段落まるごとの挿入・削除で示す。

使い方:
  python3 make_tracked.py
    -> cipherflute_wiss2026_v1.2_tracked.docx
"""
import os, copy, difflib, datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
NEW = os.path.join(HERE, "cipherflute_wiss2026_v1.2.docx")
PREV = os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_prev.docx")
OUT = os.path.join(HERE, "cipherflute_wiss2026_v1.2_tracked.docx")

AUTHOR = "CipherFlute generator"
DATE = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

_id = [1000]
def next_id():
    _id[0] += 1
    return str(_id[0])

def make_run(text, rpr_src=None, kind="normal"):
    """w:r 要素を作る。kind: normal / ins-text / del-text（del は w:delText）。"""
    r = OxmlElement("w:r")
    if rpr_src is not None:
        r.append(copy.deepcopy(rpr_src))
    t = OxmlElement("w:delText" if kind == "del-text" else "w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    return r

def wrap(kind, child_runs):
    """w:ins / w:del 要素で複数の w:r を包む。"""
    el = OxmlElement("w:ins" if kind == "ins" else "w:del")
    el.set(qn("w:id"), next_id())
    el.set(qn("w:author"), AUTHOR)
    el.set(qn("w:date"), DATE)
    for r in child_runs:
        el.append(r)
    return el

def first_rpr(p):
    for r in p._p.findall(qn("w:r")):
        rpr = r.find(qn("w:rPr"))
        if rpr is not None:
            return rpr
    return None

def clear_runs(p):
    for r in p._p.findall(qn("w:r")):
        p._p.remove(r)
    for tag in ("w:ins", "w:del"):
        for e in p._p.findall(qn(tag)):
            p._p.remove(e)

def append_child(p, el):
    p._p.append(el)

def mark_whole_insert(p):
    """段落の既存 run をすべて挿入（w:ins）として包む。"""
    runs = p._p.findall(qn("w:r"))
    if not runs:
        return
    for r in runs:
        p._p.remove(r)
    append_child(p, wrap("ins", runs))

def build_del_paragraph(new_para_for_style, old_text):
    """削除された旧段落を表す w:p 要素を作る（run を w:del で包む）。"""
    p_el = OxmlElement("w:p")
    # 近傍の新段落から段落プロパティ（スタイル）を借りる
    src_ppr = new_para_for_style._p.find(qn("w:pPr")) if new_para_for_style is not None else None
    if src_ppr is not None:
        p_el.append(copy.deepcopy(src_ppr))
    rpr_src = first_rpr(new_para_for_style) if new_para_for_style is not None else None
    r = make_run(old_text, rpr_src, kind="del-text")
    p_el.append(wrap("del", [r]))
    return p_el

def char_diff_inline(p, old_text, new_text):
    """1対1で書き換えられた段落を，文字単位の ins/del で組み直す。"""
    rpr_src = first_rpr(p)
    ppr = p._p.find(qn("w:pPr"))
    clear_runs(p)
    sm = difflib.SequenceMatcher(a=old_text, b=new_text, autojunk=False)
    def add(el):
        p._p.append(el)
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            add(make_run(new_text[j1:j2], rpr_src))
        elif tag == "insert":
            add(wrap("ins", [make_run(new_text[j1:j2], rpr_src)]))
        elif tag == "delete":
            add(wrap("del", [make_run(old_text[i1:i2], rpr_src, kind="del-text")]))
        elif tag == "replace":
            add(wrap("del", [make_run(old_text[i1:i2], rpr_src, kind="del-text")]))
            add(wrap("ins", [make_run(new_text[j1:j2], rpr_src)]))

def main():
    if not os.path.exists(PREV):
        raise SystemExit("no previous docx to compare: " + PREV)
    new_doc = Document(NEW)
    prev_doc = Document(PREV)

    # 本文テキストを持つ段落だけを対象にする（図・空段落・区切りは対象外）。
    new_paras = [p for p in new_doc.paragraphs if p.text.strip()]
    prev_paras = [p for p in prev_doc.paragraphs if p.text.strip()]
    new_txt = [p.text for p in new_paras]
    prev_txt = [p.text for p in prev_paras]

    sm = difflib.SequenceMatcher(a=prev_txt, b=new_txt, autojunk=False)
    n_ins = n_del = n_mod = 0

    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            continue
        if tag == "insert":
            for j in range(j1, j2):
                mark_whole_insert(new_paras[j])
                n_ins += 1
        elif tag == "delete":
            # 挿入位置: 新側 j1 の段落の直前。末尾なら最後の段落の直後。
            anchor = new_paras[j1] if j1 < len(new_paras) else new_paras[-1]
            for i in range(i1, i2):
                del_p = build_del_paragraph(anchor, prev_txt[i])
                if j1 < len(new_paras):
                    anchor._p.addprevious(del_p)
                else:
                    anchor._p.addnext(del_p)
                n_del += 1
        elif tag == "replace":
            if (i2 - i1) == 1 and (j2 - j1) == 1:
                char_diff_inline(new_paras[j1], prev_txt[i1], new_txt[j1])
                n_mod += 1
            else:
                anchor = new_paras[j1]
                for i in range(i1, i2):
                    del_p = build_del_paragraph(anchor, prev_txt[i])
                    anchor._p.addprevious(del_p)
                    n_del += 1
                for j in range(j1, j2):
                    mark_whole_insert(new_paras[j])
                    n_ins += 1

    new_doc.save(OUT)
    print("saved:", OUT)
    print("inserted paragraphs:", n_ins,
          "| deleted paragraphs:", n_del,
          "| in-place reworded:", n_mod)

if __name__ == "__main__":
    main()
