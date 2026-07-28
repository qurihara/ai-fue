# -*- coding: utf-8 -*-
"""印刷の向きの自由度（2026-07-27の実測）を cipherflute_wiss2026_v1.2.docx へ反映する。

docx には生成器 make_paper_wiss_v12.py に存在しない栗原さんの手編集が入っているため、
全体を再生成すると手編集が失われる。そこで python-docx で該当箇所だけを外科的に編集する。

行うこと:
 1. 現行 docx を prev/ へ退避する（make_tracked.py が「変更前」として使う）。
 2. 5章の冒頭を「3つの結果」から「4つの結果」へ直し、印刷の向きを一覧に加える。
 3. 5章の末尾（6章の見出しの直前）に、新しい節「印刷の向きの自由度」を挿入する。
 4. 3章「造形と偽装」に、向きを選べることへの言及を1文足す。
段落の書式は、既存の見出し段落と本文段落のXMLを複製して引き継ぐ。
"""
import copy
import os
import shutil

from docx import Document
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "cipherflute_wiss2026_v1.2.docx")
PREV = os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_prev.docx")

CH5_INTRO_OLD = "本章では，現時点で実機による実測が済んでいる3つの結果を報告する．"
CH5_INTRO_NEW = (
    "本章では，現時点で実機による実測が済んでいる4つの結果を報告する．すなわち，較正と音域の確定，"
    "気温に対する頑健性，印刷から復元までの通しの実証，および印刷の向きの自由度である．"
    "未実施の評価は，7章に今後の課題として整理する．")

SEC_HEAD = "印刷の向きの自由度"
SEC_BODY = [
    "日用品に埋め込むには，ホストの形に合わせて笛の向きを選べることが望ましい．そこで，音の出る窓が"
    "真上を向く従来の向きを0度とし，笛を寝かせたまま長軸まわりに回した角度を変えて，造形と発音を"
    "比べた．同一の土台の右側面，左側面，上面に，同じ音（C7）の笛を1本ずつ接した試験片を1回で印刷し，"
    "上面の1本を従来の向きの対照とした．3本が同じ印刷条件で得られるため，全部が鳴らなければ向き以外に"
    "原因があり，対照だけが鳴れば向きの問題であると切り分けられる．",

    "結果として，0度と±90度の3本すべてが鳴った．基音の平均は，対照が2180.0 Hz，窓が水平に開く2本が"
    "2177.0 Hzと2180.0 Hzであり，対照との差はそれぞれ−2.4セントと0セントであった．対照そのものも"
    "2回の吹鳴で9.5セントばらついており，向きによる差はこのばらつきに埋もれる．スロットの刻みは"
    "半音（100セント）であるから，読み取りには影響しない大きさである．",

    "一方，窓が水平に開く2本は，わずかに強い息を要し，音がやや濁って聞こえた．スペクトルを見ると，"
    "第3のピークが，対照では基音の2.99倍から3.00倍であるのに対し，水平の2本では3.09倍から3.13倍と，"
    "本来の3倍より50セントから70セント高い側へずれていた．基音と第2のピークは正常であり，濁りは"
    "高次のモードだけが整数比から外れることに由来すると考えられる．復号が読むのは基音だけであるため，"
    "この非整数化は読み取りに影響しない．",

    "したがって，埋め込みに使える向きは，窓が上，右，左の3方向に広がった．窓と吹き込み口が外気に"
    "開いてさえいれば，ホストの都合に合わせて笛を90度ずつ回してよい．窓が下を向く向きや中間の角度に"
    "ついては，同じ設計で角度を掃引する試験片を用意しており，検証を進めている．",
]

CH3_ANCHOR = "笛は日用品に融合でき，息の吹き込み口と音の出る窓だけが外部に見える形となる．"
CH3_ADD = ("笛を寝かせたまま長軸まわりに回しても音の高さが変わらないことは実測しており（5章），"
           "ホストの形に合わせて窓の向きを選べる．")


def set_text(p, txt):
    """段落の書式を保ったまま本文だけ差し替える（先頭のrunを再利用する）。"""
    for r in list(p.runs)[1:]:
        r._element.getparent().remove(r._element)
    if p.runs:
        p.runs[0].text = txt
    else:
        p.add_run(txt)


def insert_before(target, model, txt):
    """model段落のXMLを複製し，target段落の直前に本文txtで挿入する。"""
    el = copy.deepcopy(model._p)
    target._p.addprevious(el)
    p = Paragraph(el, target._parent)
    set_text(p, txt)
    return p


def main():
    os.makedirs(os.path.join(HERE, "prev"), exist_ok=True)
    shutil.copy(SRC, PREV)

    d = Document(SRC)
    ps = d.paragraphs

    # (2) 5章の冒頭
    intro = next(p for p in ps if p.text.startswith(CH5_INTRO_OLD))
    set_text(intro, CH5_INTRO_NEW)

    # (3) 6章の見出しの直前へ新しい節を挿入
    target = next(p for p in ps if p.text.strip() == "ユースケースシナリオ")
    model_h2 = next(p for p in ps if p.text.strip() == "印刷から復元までの通しの実証")
    model_body = next(p for p in ps if p.text.startswith("さらに，同じ「2026」の笛をH2D"))
    insert_before(target, model_h2, SEC_HEAD)
    for t in SEC_BODY:
        insert_before(target, model_body, t)

    # (4) 3章「造形と偽装」へ1文足す
    ch3 = next(p for p in ps if CH3_ANCHOR in p.text)
    set_text(ch3, ch3.text.replace(CH3_ANCHOR, CH3_ANCHOR + CH3_ADD))

    d.save(SRC)

    # 検証
    d2 = Document(SRC)
    texts = [p.text for p in d2.paragraphs]
    print("段落数 %d → %d（+%d）" % (len(ps), len(texts), len(texts) - len(ps)))
    print("5章冒頭が4つに:", any(t.startswith("本章では，現時点で実機による実測が済んでいる4つの結果") for t in texts))
    print("新しい節の見出し:", texts.count(SEC_HEAD), "個")
    print("新しい節の本文:", sum(1 for t in SEC_BODY if t in texts), "/", len(SEC_BODY))
    print("3章への追記:", any(CH3_ADD in t for t in texts))
    print("概要は無改変:", texts[4].startswith("概要．"))


if __name__ == "__main__":
    main()
