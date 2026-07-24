# -*- coding: utf-8 -*-
"""cipherflute_wiss2026_v1.2.docx への編集を適用する。

手順:
 1. 原本(栗原さんの編集履歴つき)を prev/ にバックアップ。
 2. 栗原さんの変更履歴を accept(=w:ins を昇格・w:del を除去)して確定テキスト版を作り、
    prev/cipherflute_wiss2026_v1.2_prev.docx に保存(make_tracked.py の比較元)。
 3. 確定版に本セッションの修正を適用:
    - カタカナ「ビット」→bit,「バイト」→byte に統一(「ビットコイン」は固有名詞で除外)。
    - 4章のAUC予告文(発見率で定量化する予定)を削除。
    - 3章(設計)の較正定数 A=86718,e=-11.81 の記載を削除(5章に残す)。
    - デコーダ節にデコード用Webアプリの本文と図5(復号ソフト画面)の参照・キャプションを追加。
 4. cipherflute_wiss2026_v1.2.docx を上書き保存。
"""
import os, re, shutil, zipfile
from lxml import etree
from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "cipherflute_wiss2026_v1.2.docx")
NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
W = '{%s}' % NS['w']
os.makedirs(os.path.join(HERE, "prev"), exist_ok=True)

# 1. 原本バックアップ
shutil.copy(SRC, os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_kurihara_tracked.docx"))

# 2. accept changes → 確定版
zin = zipfile.ZipFile(SRC)
root = etree.fromstring(zin.read('word/document.xml'))
for el in list(root.iter(W + 'del')):        # 削除履歴は本文から除く
    el.getparent().remove(el)
for ins in list(root.iter(W + 'ins')):        # 挿入履歴は確定文へ昇格
    par = ins.getparent()
    idx = list(par).index(ins)
    for ch in list(ins):
        par.insert(idx, ch); idx += 1
    par.remove(ins)
ACC = os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_prev.docx")
newxml = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
zout = zipfile.ZipFile(ACC, 'w', zipfile.ZIP_DEFLATED)
for item in zin.namelist():
    zout.writestr(item, newxml if item == 'word/document.xml' else zin.read(item))
zout.close()
zin.close()

# 3. 確定版に修正を適用
d = Document(ACC)


def set_text(p, txt):
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.add_run(txt)


def bitbyte(s):
    s = s.replace("ビットコイン", "\x00BC\x00")
    s = s.replace("ビット", "bit").replace("バイト", "byte")
    return s.replace("\x00BC\x00", "ビットコイン")


n_bit = 0
deco_idx = None
for i, p in enumerate(d.paragraphs):
    t = p.text
    if not t.strip():
        continue
    new = t
    # (a) bit/byte 統一
    if "ビット" in new or "バイト" in new:
        new = bitbyte(new)
    # (b) 4章 AUC予告の削除
    if "発見率の調査で定量化する予定である" in new:
        new = new.replace("この点は今後，発見率の調査で定量化する予定である（7章）．", "")
        new = new.replace("この点は今後，発見率の調査で定量化する予定である（7章）。", "")
    # (c) 3章(設計)の較正定数の削除。5章に残す。
    if "A = 86718" in new and "詳細は5章" in new:
        new = re.sub(r"A = 86718，e = −11\.81 であり13本すべてが狙いから±6セント以内で鳴った（詳細は5章）．", "", new)
        new = re.sub(r"A = 86718，e = −11\.81 であり13本すべてが狙いから±6セント以内で鳴った（詳細は5章）。", "", new)
    if new != t:
        set_text(p, new)
        n_bit += 1
    # デコーダ節見出しの位置を控える
    if t.strip().startswith("秘密分散との組み合わせと，正しさの確かめ方"):
        deco_idx = i

# (d) デコーダ本文＋図5を、デコーダ節見出しの後の空段落へ入れる
DECO_BODY = ("復元した値の正しさの確認とは別に，本手法にはデコード用のWebアプリケーションを用意した．"
             "スマートフォンのマイクへの入力を周波数分析するだけで動作し，ネットワーク接続を必要としない．"
             "利用者は器物を吹くだけで，画面に読み取った音の並びと復元結果が表示される．図5にその利用の様子を示す．")
FIG5_CAP = "図5　復号ソフトの画面．吹いた音の周波数を解析し，スロット列と復元した値を表示する．"
if deco_idx is not None:
    # 見出しの直後にある空段落を2つ使う（本文・図5キャプション）
    empties = []
    for j in range(deco_idx + 1, min(deco_idx + 8, len(d.paragraphs))):
        if not d.paragraphs[j].text.strip():
            empties.append(j)
    # 秘密分散本文の後の空段落を優先（deco_idx+2, +3 付近）
    if len(empties) >= 2:
        set_text(d.paragraphs[empties[0]], DECO_BODY)
        set_text(d.paragraphs[empties[1]], FIG5_CAP)
        placed = True
    else:
        placed = False
else:
    placed = False

d.save(SRC)

# 検証
d2 = Document(SRC)
rem_bit = sum(p.text.count("ビット") - p.text.count("ビットコイン") for p in d2.paragraphs)
rem_byte = sum(p.text.count("バイト") for p in d2.paragraphs)
has_auc = any("発見率の調査で定量化する予定" in p.text for p in d2.paragraphs)
calib_ch3 = [p.text[:40] for p in d2.paragraphs if "A = 86718" in p.text]
has_deco = any("デコード用のWebアプリ" in p.text for p in d2.paragraphs)
print("編集した段落数:", n_bit)
print("残カタカナ bit=%d byte=%d (0が目標)" % (rem_bit, rem_byte))
print("AUC予告 残存:", has_auc, " (Falseが目標)")
print("A=86718 を含む段落数:", len(calib_ch3), "→", calib_ch3)
print("デコーダ本文 追加:", has_deco, " 図5配置:", placed)
