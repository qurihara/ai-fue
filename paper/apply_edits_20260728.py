# -*- coding: utf-8 -*-
"""2026-07-28 の結果を論文へ反映する（外科的な追記のみ）。

この日に得た、論文に載せる価値のある結果は次の3つである。

  1. 12音を各3本、計36本の較正コームを実測し、吹き方のばらつきと造形のばらつきを
     分けて評価できた。12音すべてが1つの式 f = A/(L+e) に乗る。
  2. 128bitの秘密を、スプール2枚に分けて埋め込む実装ができた（1枚だけでは復元できない）。
     ユースケースの節が「想定」から「作成中」へ進む。
  3. 外見を統一する設計では、外形の長さを最長管ちょうどにすると最低音のボアだけが
     短くなる。これは同じ方式を採る人が必ず踏む設計上の制約なので、1文だけ実装の
     注意として書く（不具合の経緯としては書かない）。

docx には生成器に無い栗原さんの手編集があるので、全体再生成はせず該当箇所だけを直す。
"""
import copy
import os
import shutil

from docx import Document

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "cipherflute_wiss2026_v1.2.docx")
PREV = os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_prev_20260728.docx")

# 1. 較正と音域の確定の節に、36本コームの評価を1段落足す
ANCHOR_CALIB = "厚壁ヘッド（mini10）の笛を用い，狙った音の高さに合わせ込む較正を行った．"
NEW_CALIB = (
    "さらに，外見を統一した笛（外形を最長管に揃え，内部の仕切り壁で管の長さを決める形）"
    "について，12音（G#6からG7）を各3本，計36本並べた較正用のコームを印刷し，"
    "3回通して吹いて評価した．低い4音は1本ずつ吹き直して位置を確定させている．"
    "同じ笛を繰り返し吹いたときのばらつきは15.2セントであった．同じ音の3本のあいだの"
    "散らばりは11.0セントで，これは吹き方のばらつきだけで説明がつく大きさである．"
    "したがって造形そのもののばらつきは，この測り方では吹き方のばらつきに埋もれており，"
    "多くとも11セント程度と言える．音の高さの置き場を半音（100セント）刻みとする現在の"
    "設計に対しては，十分に小さい．また，実測した管の長さを用いて f = A/(L+e) を"
    "当てはめ直すと，12音すべてが1つの式に乗り，残差は二乗平均平方根で11.3セントに"
    "収まった。音によって系統的にずれる成分は残らず，最も高いG7も3本とも±20セント以内"
    "であったことから，G#6からG7までの12スロットを使える見通しが得られた．"
)

# 2. 外見統一の設計上の注意（造形と偽装の節の末尾へ1文）
ANCHOR_UNIFORM = "笛の素材について，食品接触に適合していて扱いが容易なBambu Lab PLA Pureで"
NEW_UNIFORM = (
    "なお，外見を統一するために外形の長さを最長管に揃える場合，外形の長さを最長管"
    "ちょうどにしてはならない．最も低い音の笛だけは内部に仕切り壁を置く余地がなくなり，"
    "閉じた端の壁の厚みのぶんだけ管が短くなって，音が高い側へずれるためである．"
    "実装では最長管に数ミリの余白を足した長さを外形とし，どの音にも仕切り壁が入るようにした．"
)

# 3. 128bitのユースケースを「想定」から「作成中」へ
OLD_HEAD_128 = "128bitの復元用パスフレーズの保管（想定）"
NEW_HEAD_128 = "128bitの復元用パスフレーズの保管（作成中）"
OLD_BODY_128 = "12単語のシード相当の128bitを，多数の笛としてシンプルな箱の底面に埋め込む使い方を想定している．"
NEW_BODY_128 = (
    "12単語のシード相当の128bitを，多数の笛として日用品に埋め込む使い方である．"
    "隣り合う笛が同じ音にならない符号化と12スロットの組み合わせでは，128bitは笛49本"
    "（基準笛1本を含む）になる．これを，フィラメントを巻くスプールの円盤2枚へ25本と24本に"
    "分けて埋め込んだものを作成中である．1つの符号語を2枚に割っているため，"
    "片方だけでは復元できない．保管場所を分けるという運用が，物のかたちでそのまま実現される．"
    "1枚の円盤に置ける本数には限りがあり，笛の外形と円盤の半径から26本が上限である"
    "（27本で隣どうしの隙間が1 mmを切る）．"
)


def paragraph_after(doc, anchor_head):
    for i, p in enumerate(doc.paragraphs):
        if p.text.startswith(anchor_head):
            return i, p
    raise SystemExit("見つからない: " + anchor_head[:30])


def insert_after(paragraph, text):
    """同じ書式の段落を直後に足す。"""
    new = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new, paragraph._parent)
    for r in para.runs[1:]:
        r._r.getparent().remove(r._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)
    return para


def main():
    os.makedirs(os.path.dirname(PREV), exist_ok=True)
    shutil.copy2(SRC, PREV)
    doc = Document(SRC)

    _, p_cal = paragraph_after(doc, ANCHOR_CALIB[:30])
    insert_after(p_cal, NEW_CALIB)
    print("較正の節に36本コームの評価を足した")

    _, p_uni = paragraph_after(doc, ANCHOR_UNIFORM[:20])
    p_uni.add_run(NEW_UNIFORM)
    print("造形と偽装の節に外形の余白の注意を足した")

    hit = 0
    for p in doc.paragraphs:
        if p.text.strip() == OLD_HEAD_128:
            for r in p.runs[1:]:
                r._r.getparent().remove(r._r)
            p.runs[0].text = NEW_HEAD_128
            hit += 1
        elif p.text.startswith(OLD_BODY_128[:25]):
            rest = p.text[len(OLD_BODY_128):]
            for r in p.runs[1:]:
                r._r.getparent().remove(r._r)
            p.runs[0].text = NEW_BODY_128 + rest
            hit += 1
    if hit != 2:
        raise SystemExit("128bitの節の書き換えに失敗した（%d箇所）" % hit)
    print("128bitのユースケースを作成中の実装として書き直した")

    doc.save(SRC)
    print("保存した:", os.path.relpath(SRC, HERE))


if __name__ == "__main__":
    main()
