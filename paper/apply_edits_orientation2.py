# -*- coding: utf-8 -*-
"""印刷の向きの節を、角度掃引と縦置きの結果まで含めて書き直す。

2026-07-27に「窓が上・右・左の3方向」まで書いたが、7月28日に角度を掃引して
[* -135度から+135度まで連続して鳴る]（360度のうち270度）と確定した。あわせて
縦置きの結果も分かったので、節の3段落目と4段落目を差し替え、1段落足す。

docx には生成器に無い栗原さんの手編集があるので、全体再生成はせず該当箇所だけを直す。
"""
import copy
import os
import shutil

from docx import Document
from docx.text.paragraph import Paragraph

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "cipherflute_wiss2026_v1.2.docx")
PREV = os.path.join(HERE, "prev", "cipherflute_wiss2026_v1.2_prev.docx")

OLD3_HEAD = "一方，窓が水平に開く2本は，わずかに強い息を要し"
NEW3 = (
    "一方，窓が真上を向かない向きでは，わずかに強い息を要し，音がやや濁って聞こえた．"
    "スペクトルを見ると，第3のピークが，対照では基音の2.99倍から3.00倍であるのに対し，"
    "窓が水平の2本では3.09倍から3.13倍と，本来の3倍より50セントから70セント高い側へ"
    "ずれていた．基音と第2のピークは正常であり，濁りは高次のモードだけが整数比から"
    "外れることに由来すると考えられる．復号が読むのは基音だけであるため，この非整数化は"
    "読み取りに影響しない．")

OLD4_HEAD = "したがって，埋め込みに使える向きは，窓が上，右，左の3方向に広がった．"
NEW4 = (
    "続いて，同じ設計で角度を掃引した試験片を印刷し，0度から180度までを確かめた．"
    "その結果，0度，±45度，±90度，±135度はいずれも鳴り，180度（窓が真下）だけが"
    "鳴らなかった．180度では管の天井が正しく被覆されず，窓の造形も崩れていた．"
    "したがって，埋め込みに使える向きは −135度から +135度までの連続した範囲であり，"
    "360度のうち270度を占める．使えないのは真下まわりの90度だけである．"
    "日用品への埋め込みにおいて，笛の向きに関する制約はほぼ無くなったといえる．")

NEW5 = (
    "笛の長軸を立てる置き方も試した．この場合，吹き込み口から音の出る窓へ至る風道は"
    "高さ約1 mmの垂直な細い管になる．吹き込み口を上に向けるとこの風道が埋まって"
    "息が入らず，鳴らなかった．吹き込み口を下に向けた場合は鳴ったが，低い音は狙いより"
    "95セント高く，高い音では一回の吹鳴で複数の基音が同時に現れて基音が定まらなかった．"
    "注意すべきは，聴いた印象と測定が食い違ったことである．耳では高い音がよく鳴り"
    "低い音はかすれて聞こえたが，測定では低い音の方が安定していた．復号器が読むのは"
    "基音ひとつであるから，鳴っているかどうかを聴感で判断してはならない．"
    "以上より，長軸を立てる置き方は現時点では推奨しない．")


def set_text(p, txt):
    for r in list(p.runs)[1:]:
        r._element.getparent().remove(r._element)
    if p.runs:
        p.runs[0].text = txt
    else:
        p.add_run(txt)


def main():
    os.makedirs(os.path.join(HERE, "prev"), exist_ok=True)
    shutil.copy(SRC, PREV)
    d = Document(SRC)
    ps = d.paragraphs

    p3 = next(p for p in ps if p.text.startswith(OLD3_HEAD))
    set_text(p3, NEW3)
    p4 = next(p for p in ps if p.text.startswith(OLD4_HEAD))
    set_text(p4, NEW4)

    el = copy.deepcopy(p4._p)
    p4._p.addnext(el)
    set_text(Paragraph(el, p4._parent), NEW5)

    d.save(SRC)

    d2 = Document(SRC)
    t = [q.text for q in d2.paragraphs]
    print("段落数 %d → %d（+%d）" % (len(ps), len(t), len(t) - len(ps)))
    print("270度の記述:", any("360度のうち270度" in x for x in t))
    print("縦置きの記述:", any("長軸を立てる置き方" in x for x in t))
    print("概要は無改変:", t[4].startswith("概要．"))


if __name__ == "__main__":
    main()
